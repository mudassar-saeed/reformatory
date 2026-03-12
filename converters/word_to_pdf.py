import os
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY


def convert_word_to_pdf(input_path: str, output_path: str) -> None:
    """
    Converts a Word (.docx) file to a PDF file.

    :param input_path: Path to the input DOCX file
    :param output_path: Path where the output PDF will be saved
    :raises FileNotFoundError: If input file does not exist
    :raises Exception: If conversion fails
    """

    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")

    try:
        doc = Document(input_path)

        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        base_styles = getSampleStyleSheet()

        style_normal = ParagraphStyle(
            "CustomNormal",
            parent=base_styles["Normal"],
            fontSize=11,
            leading=16,
            spaceAfter=4,
        )

        style_h1 = ParagraphStyle(
            "CustomH1",
            parent=base_styles["Heading1"],
            fontSize=18,
            leading=24,
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor("#1a1a1a"),
        )

        style_h2 = ParagraphStyle(
            "CustomH2",
            parent=base_styles["Heading2"],
            fontSize=15,
            leading=20,
            spaceBefore=10,
            spaceAfter=4,
            textColor=colors.HexColor("#2a2a2a"),
        )

        style_h3 = ParagraphStyle(
            "CustomH3",
            parent=base_styles["Heading3"],
            fontSize=13,
            leading=18,
            spaceBefore=8,
            spaceAfter=4,
            textColor=colors.HexColor("#3a3a3a"),
        )

        story = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Handle paragraphs
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break

                if para is None:
                    continue

                text = para.text.strip()
                style_name = para.style.name if para.style else "Normal"

                if not text:
                    story.append(Spacer(1, 6))
                    continue

                # Escape special XML chars
                text = (text
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;"))

                # Pick style based on heading level
                if "Heading 1" in style_name:
                    p_style = style_h1
                elif "Heading 2" in style_name:
                    p_style = style_h2
                elif "Heading 3" in style_name:
                    p_style = style_h3
                else:
                    p_style = style_normal

                # Apply bold/italic from runs
                parts = []
                for run in para.runs:
                    run_text = (run.text
                                .replace("&", "&amp;")
                                .replace("<", "&lt;")
                                .replace(">", "&gt;"))
                    if not run_text:
                        continue
                    if run.bold and run.italic:
                        parts.append(f"<b><i>{run_text}</i></b>")
                    elif run.bold:
                        parts.append(f"<b>{run_text}</b>")
                    elif run.italic:
                        parts.append(f"<i>{run_text}</i>")
                    else:
                        parts.append(run_text)

                formatted = "".join(parts) if parts else text
                story.append(Paragraph(formatted, p_style))

            elif tag == "tbl":
                # Handle tables
                for tbl in doc.tables:
                    if tbl._element is element:
                        data = []
                        for row in tbl.rows:
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                cell_text = (cell_text
                                             .replace("&", "&amp;")
                                             .replace("<", "&lt;")
                                             .replace(">", "&gt;"))
                                row_data.append(Paragraph(cell_text, style_normal))
                            data.append(row_data)

                        if data:
                            table = Table(data, repeatRows=1)
                            table.setStyle(TableStyle([
                                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4f46e5")),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                                ("FONTSIZE", (0, 0), (-1, -1), 10),
                                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                                 [colors.white, colors.HexColor("#f3f4f6")]),
                                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
                                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                                ("TOPPADDING", (0, 0), (-1, -1), 4),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                            ]))
                            story.append(Spacer(1, 8))
                            story.append(table)
                            story.append(Spacer(1, 8))
                        break

        if not story:
            story.append(Paragraph("(Empty document)", style_normal))

        pdf_doc.build(story)

    except Exception as e:
        raise Exception(f"Word to PDF conversion failed: {str(e)}")
